"""Convert PRICING_AND_TIERS.md to a polished Word report.

Targets python-docx 1.2.0. Handles the markdown constructs actually used in the
source document: ATX headings, GitHub tables (with alignment), bullet/numbered
lists, fenced code blocks, blockquotes, horizontal rules, inline bold/italic,
inline code, and hyperlinks.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path(__file__).resolve().parents[1] / "PRICING_AND_TIERS.md"
DST = Path(__file__).resolve().parents[1] / "LexDraft_Pricing_and_Tiers.docx"


# ---------- styling helpers ----------

ACCENT = RGBColor(0x1F, 0x2A, 0x44)        # deep navy
MUTED = RGBColor(0x55, 0x5F, 0x6D)
HEADER_FILL = "1F2A44"
ZEBRA_FILL = "F4F5F7"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D0D5DD")
        borders.append(b)
    tc_pr.append(borders)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F6FEB")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    new_run.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------- inline parsing ----------

INLINE_PATTERN = re.compile(
    r"""(
        \*\*(?P<bold>[^*]+)\*\* |
        \[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\) |
        `(?P<code>[^`]+)` |
        \*(?P<italic>[^*]+)\*
    )""",
    re.VERBOSE,
)


def render_inline(paragraph, text: str, base_size: Pt | None = None) -> None:
    """Render markdown inline formatting into a paragraph."""
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            if base_size:
                run.font.size = base_size
        if match.group("bold"):
            run = paragraph.add_run(match.group("bold"))
            run.bold = True
            if base_size:
                run.font.size = base_size
        elif match.group("link_text"):
            add_hyperlink(paragraph, match.group("link_url"), match.group("link_text"))
        elif match.group("code"):
            run = paragraph.add_run(match.group("code"))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif match.group("italic"):
            run = paragraph.add_run(match.group("italic"))
            run.italic = True
            if base_size:
                run.font.size = base_size
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if base_size:
            run.font.size = base_size


# ---------- block-level helpers ----------

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def parse_alignment(separator_row: list[str]) -> list[str]:
    aligns = []
    for cell in separator_row:
        cell = cell.strip()
        left = cell.startswith(":")
        right = cell.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(doc: Document, rows: list[list[str]], aligns: list[str]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for r_idx, row in enumerate(rows):
        is_header = r_idx == 0
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            if is_header:
                set_cell_shading(cell, HEADER_FILL)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, ZEBRA_FILL)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = ALIGN_MAP[aligns[c_idx] if c_idx < len(aligns) else "left"]
            value = row[c_idx] if c_idx < len(row) else ""
            render_inline(para, value)
            for run in para.runs:
                if is_header:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                run.font.size = Pt(9.5) if not is_header else Pt(10)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    para.paragraph_format.space_after = Pt(6)
    if level == 1:
        para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    run.font.size = Pt(sizes.get(level, 11))


def add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    render_inline(para, text)


def add_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    render_inline(para, text)


def add_numbered(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(2)
    render_inline(para, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F5F7")
    pPr.append(shd)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        if i < len(lines) - 1:
            para.add_run().add_break()


def add_blockquote(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(8)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "1F2A44")
    pBdr.append(left)
    pPr.append(pBdr)
    render_inline(para, text)
    for run in para.runs:
        run.font.color.rgb = MUTED
        run.italic = True


def add_horizontal_rule(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D5DD")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------- top-level parser ----------

def parse(doc: Document, lines: list[str]) -> None:
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # fenced code blocks
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # tables — header row then separator row
        if stripped.startswith("|") and i + 1 < len(lines):
            sep = lines[i + 1].strip()
            if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", sep):
                header = split_table_row(stripped)
                aligns = parse_alignment(split_table_row(sep))
                rows = [header]
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    rows.append(split_table_row(lines[j].strip()))
                    j += 1
                add_table(doc, rows, aligns)
                i = j
                continue

        # blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            add_blockquote(doc, text)
            i += 1
            continue

        # bullet list
        if re.match(r"^[-*]\s+", stripped):
            add_bullet(doc, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", stripped):
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # plain paragraph — combine continuation lines
        buf = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip("\n")
            nxt_stripped = nxt.strip()
            if (
                not nxt_stripped
                or nxt_stripped.startswith("#")
                or nxt_stripped.startswith("|")
                or nxt_stripped.startswith(">")
                or nxt_stripped.startswith("```")
                or re.match(r"^[-*]\s+", nxt_stripped)
                or re.match(r"^\d+\.\s+", nxt_stripped)
                or re.fullmatch(r"-{3,}", nxt_stripped)
            ):
                break
            buf.append(nxt_stripped)
            j += 1
        add_paragraph(doc, " ".join(buf))
        i = j


# ---------- cover ----------

def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = eyebrow.add_run("LexDraft · Product strategy")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    run.bold = True

    title = doc.add_paragraph()
    title_run = title.add_run("Pricing, Tiering & Account Strategy")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        "Solo · Practice · Firm · SuperAdmin — market research, feature gating, "
        "account-creation flows, billing, RBAC, and DPDP-aligned compliance."
    )
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(18)

    meta = doc.add_paragraph()
    for label, value in (
        ("Status", "Draft v1"),
        ("Date", "2026-05-06"),
        ("Owner", "Product"),
        ("Audience", "Engineering, Design, GTM"),
    ):
        run = meta.add_run(f"{label}:  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = ACCENT
        v = meta.add_run(f"{value}    ")
        v.font.size = Pt(10)
        v.font.color.rgb = MUTED

    add_horizontal_rule(doc)


def strip_frontmatter(text: str) -> str:
    """Drop the first H1 + the meta block we render in the cover instead."""
    lines = text.splitlines()
    # remove the H1 title line
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    # drop leading blank lines
    while lines and not lines[0].strip():
        lines = lines[1:]
    # drop the bold meta block (Status / Owner / Audience / Scope) if present
    while lines and lines[0].startswith("**"):
        lines = lines[1:]
    while lines and not lines[0].strip():
        lines = lines[1:]
    return "\n".join(lines)


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    text = strip_frontmatter(text)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc)
    parse(doc, text.splitlines())
    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
